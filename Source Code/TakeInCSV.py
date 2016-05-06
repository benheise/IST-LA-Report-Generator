
from __future__ import division
from decimal import *
from docx import Document


def conversion_table_construct():
    marker = 0
    conversion_table = {}
    grade = 0
    with open('checkwords.txt') as textids:
        for row in textids:
            if '##' in row:
                marker = 0
            elif 'Eval Words' in row:
                marker = 1
            elif 'Ignore Words' in row:
                marker = 0
            elif marker is 1 and row is not "\n" or '':
                grade += 1
                conversion_table[row[:-1]] = grade
    textids.close()
    print "looking for these words"
    print conversion_table
    return conversion_table


def avoid_table_construct():
    marker = 0
    avoid_list = []
    with open('checkwords.txt') as textids:
        for row in textids:
            if '##' in row:
                continue
            elif 'Ignore Words' in row:
                marker = 1
            elif marker is 1 and row is not "\n":
                avoid_list.append(row[:-1])
    textids.close()
    print "avoiding these words"
    print avoid_list
    return avoid_list


def convert_to_number(convert_text, conversion_table):
    number = 0
    number_array = []
    for text in convert_text:
        for key in conversion_table.keys():
            if key in text:
                number = conversion_table[key]
                number_array.append(number)
            elif number == 0 or number == 'Null':
                number = 'Null'
    return number_array


def trim_response(text_to_trim):
    return_text = []
    for i in range(10, len(text_to_trim)):
        return_text.append(text_to_trim[i])
    return return_text


def trim_question(question_to_trim):
    question_text = []
    for i in range(10, len(question_to_trim)-4):
        if '${' in question_to_trim[i]:
            question_text_p1 = question_to_trim[i].split('${', 1)[-2]
            question_text_p2a = question_to_trim[i].split('${', 1)[-1]
            question_text_p2 = question_text_p2a.split('}', 1)[-1]
            question_to_append = question_text_p1 + ' ' + question_text_p2
            if '\xc2\xa0' in question_to_append:
                question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
            elif 'Assistant\xe2\x80\x99s' in question_to_append:

                question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                question_to_append = question_text_p5 + "Assistant's" + question_text_p6
            question_text.append(question_to_append)
        else:
            question_text.append(question_to_trim[i])
    return question_text


def response_find(name, question_to_search, responses):
    ans = []
    for entry in responses:
        if name in entry:
            ans.append(entry[question_to_search])
    return ans


def create_individual_report(response_list, name):
    report_data = []
    for entry in response_list:
        if name in entry[1]:
            report_data.append(entry)
    return report_data


def create_name_list(response_list):
    name_list = []
    for entry in response_list:
        if entry[1] in name_list:
            continue
        else:
            name_list.append(entry[1])
    return name_list


def response_options(question_number, report_array):
    rating_list = []
    for entry in report_array:
        add_rating = entry[question_number]
        if add_rating not in rating_list:
            rating_list.append(add_rating)
    return rating_list


def fill_table(question_number, document, table_number, report_container):
    rating_list = []
    for entry in report_container:
        add_rating = entry[question_number]
        rating_list.append(add_rating)
    rating_catagories = response_options(question_number, report_container)
    rating_catagories.sort()
    rating_list.sort()
    total = 0
    new_row = document.tables[table_number]
    toi = document.tables[table_number]

    for rating in rating_catagories:
        new_row.add_row()
        toi.cell(-1, 0).text = rating
        for entry in rating_list:
            if entry == rating:
                total += 1
            else:
                total = total
        getcontext().prec = 2
        value = int(document.tables[0].cell(1, 1).text)
        average = Decimal(total/int(document.tables[0].cell(1, 1).text))
        print str(total) + "/" + str(value) + " = " + str(average)
        toi.cell(-1, 1).text = str(total)
        toi.cell(-1, 2).text = str(average*100) + '%'
        total = 0


def print_individual_report(report_holder, questions_list):
    document = Document('Individual Template.docx')
    document.paragraphs[0].add_run(report_holder[0][1])
    document.tables[0].cell(1, 1).text = str(len(report_holder))
    classes = ''

    for entry in report_holder:
        add_class = entry[0]
        if add_class not in classes:
            classes = classes + '\n' + add_class
    document.paragraphs[1].add_run(classes)
    for i in range(1, len(questions_list)-1):
        fill_table(i+1, document, i, report_holder)
        question_to_print = ''.join([j if ord(j) < 128 else ' ' for j in questions_list[i+1]])
        document.paragraphs[i+1].add_run(question_to_print)
    name = report_holder[0][1] + '.docx'
    if name == '#N/A.docx':
        name = report_holder[0][0] + '.docx'
        document.save('reports/'+name)
    else:
        document.save('reports/'+name)


def calc_average_for_array(array):
    holder = 0
    for number in array:
        holder = holder + number
    average = holder / len(array)
    return average


def print_master_report(reports, convert_table):
    q1_average = []
    q2_average = []
    q4_average = []
    q5_average = []
    q3_average = []
    document = Document('Master Template.docx')
    document.paragraphs[1].add_run(str(len(reports)))

    for part in reports:
        for entry in part:
            q1_average.append(entry[2])
            q2_average.append(entry[3])
            q3_average.append(entry[4])
            q4_average.append(entry[5])
            q5_average.append(entry[6])
    q1_average = calc_average_for_array(convert_to_number(q1_average, convert_table))
    q2_average = calc_average_for_array(convert_to_number(q2_average, convert_table))
    q3_average = calc_average_for_array(convert_to_number(q3_average, convert_table))
    q4_average = calc_average_for_array(convert_to_number(q4_average, convert_table))
    q5_average = calc_average_for_array(convert_to_number(q5_average, convert_table))
    document.paragraphs[3].add_run(str(q1_average))
    document.paragraphs[5].add_run(str(q2_average))
    document.paragraphs[7].add_run(str(q3_average))
    document.paragraphs[9].add_run(str(q4_average))
    document.paragraphs[11].add_run(str(q5_average))
    document.save('Master Report.docx')


with open("SP16 TA LA Eval.csv", "r") as readfile:
    questions = []
    response_holder = []
    report_list = []
    avoid_table = avoid_table_construct()
    number_table = conversion_table_construct()
    counter = 0

    for line in readfile:
        response = line.split(',')
        if counter is 0:
            counter += 1
        elif counter is 1:
            questions = trim_question(response)
            counter += 1
        else:
            response_holder.append(trim_response(response))
            counter += 1

    for learning_assistant in create_name_list(response_holder):
        if learning_assistant not in report_list:
            report = create_individual_report(response_holder, learning_assistant)
            report_list.append(report)
            print_individual_report(report, questions)
    print_master_report(report_list, number_table)
readfile.close()
