from __future__ import division
from PyQt4 import QtGui
from PyQt4.Qt import *
from PyQt4.QtCore import QThread
from csv import *
from collections import defaultdict
from docx import Document
from PyQt4 import QtCore, QtGui
from decimal import *
import sys, os, csv, difflib,random
import design


class Main_App(QtGui.QMainWindow, design.Ui_MainWindow):

    def __init__(self,parent=None):
        super(self.__class__, self).__init__()
        self.setupUi(self)
        self.createDirectories()

        self.filepath = ''
        self.template = Document('Template.docx')
        self.qualtricsButton.clicked.connect(self.createQuatrics)
        self.loadFileButton.clicked.connect(self.loadFile)
        self.allReportsButton.clicked.connect(self.allReports)
        self.singleReportButton.clicked.connect(self.singleReport)
        self.masterReportButton.clicked.connect(self.masterReport)

    def allReports(self):
        filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(filepath) > 0:
            self.myThread = Generate_All_Reports(filepath, self.template)
            self.connect(self.myThread, SIGNAL("finished()"), self.Done)
            self.myThread.start()
    def createQuatrics(self):
        filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(filepath) > 0:
            self.myThread = Create_Qualtrics_CSV(filepath)
            self.connect(self.myThread, SIGNAL("finished()"), self.QualDone)
            self.myThread.start()
    def QualDone(self):
        QtGui.QMessageBox.information(self, "Done!", "Qualtrics CSV Created in Current Working Directory!")

    def IndivDone(self):
        QtGui.QMessageBox.information(self, "Done!", "Check the Reports Folder for your Report")

    def Done(self):
        QtGui.QMessageBox.information(self, "Done!", "All TA/LA reports generated in Reports Folder!")

    def loadFile(self):
        self.filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(self.filepath) > 0:
            self.myThread = fillNames(self.filepath)
            self.connect(self.myThread,SIGNAL('addLaName(QString)'), self.addName)
            self.myThread.start()

    def addName(self,name):
        self.listWidget.addItem(name)

    def singleReport(self):
        currentPos = self.listWidget.currentRow()
        item = self.listWidget.item(currentPos)
        LaName = item.text()
        if len(self.filepath) > 0:
            self.myThread = generateIndividualReport(self.filepath,LaName,self.template)
            self.connect(self.myThread, SIGNAL("finished()"), self.IndivDone)
            self.myThread.start()
        else:
            QtGui.QMessageBox.information(self, "Error", "TA/LA CSV Not Loaded Yet! Please load to proceed")

    def createDirectories(self):
        if not os.path.exists('Reports'):
            os.makedirs("Reports")

    def masterReport(self):
        filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(filepath) > 0:
            self.myThread = Generate_Master_Report(filepath, self.template)
            self.connect(self.myThread, SIGNAL("finished()"), self.Done)
            self.myThread.start()

class generateIndividualReport(QThread):
    def __init__(self, filepath, La_Name, Template):
        QThread.__init__(self)
        self.filepath = filepath
        self.LaName = La_Name
        self.template = Template
    def __del__(self):
        self.wait()

    def conversion_table_construct(self):
        marker = 0
        conversion_table = {}
        grade = 0
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    marker = 0
                elif 'Eval Words' in row:
                    marker = 1
                elif 'Ignore Words' in row:
                    marker = 0
                elif marker is 1 and row is not "\n" or '':
                    grade += 1
                    conversion_table[row[:-1]] = grade
        textids.close()
        print "looking for these words"
        print conversion_table
        return conversion_table

    def avoid_table_construct(self):
        marker = 0
        avoid_list = []
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Ignore Words' in row:
                    marker = 1
                elif marker is 1 and row is not "\n":
                    avoid_list.append(row[:-1])
        textids.close()
        print "avoiding these words"
        print avoid_list
        return avoid_list

    def convert_to_number(self, convert_text, conversion_table):
        number = 0
        number_array = []
        for text in convert_text:
            for key in conversion_table.keys():
                if key in text:
                    number = conversion_table[key]
                    number_array.append(number)
                elif number == 0 or number == 'Null':
                    number = 'Null'
        return number_array

    def trim_response(self, text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self, question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:

                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

    def response_find(self, name, question_to_search, responses):
        ans = []
        for entry in responses:
            if name in entry:
                ans.append(entry[question_to_search])
        return ans

    def create_individual_report(self, response_list, name):
        report_data = []
        for entry in response_list:
            if name in entry[1]:
                report_data.append(entry)
        return report_data

    def create_name_list(self, response_list):
        name_list = []
        for entry in response_list:
            if entry[1] in name_list:
                continue
            else:
                name_list.append(entry[1])
        return name_list

    def response_options(self, question_number, report_array):
        rating_list = []
        for entry in report_array:
            add_rating = entry[question_number]
            if add_rating not in rating_list:
                rating_list.append(add_rating)
        return rating_list

    def fill_table(self, question_number, document, table_number, report_container):
        rating_list = []
        for entry in report_container:
            add_rating = entry[question_number]
            rating_list.append(add_rating)
        rating_catagories = self.response_options(question_number, report_container)
        rating_catagories.sort()
        rating_list.sort()
        total = 0
        new_row = document.tables[table_number]
        toi = document.tables[table_number]

        for rating in rating_catagories:
            new_row.add_row()
            toi.cell(-1, 0).text = rating
            for entry in rating_list:
                if entry == rating:
                    total += 1
                else:
                    total = total
            getcontext().prec = 2
            average = Decimal(total / int(document.tables[0].cell(1, 1).text))
            toi.cell(-1, 1).text = str(total)
            toi.cell(-1, 2).text = str(average * 100) + '%'
            total = 0

    def print_individual_report(self, report_holder, questions_list):
        document = Document('Individual Template.docx')
        document.paragraphs[0].add_run(report_holder[0][1])
        document.tables[0].cell(1, 1).text = str(len(report_holder))
        classes = ''

        for entry in report_holder:
            add_class = entry[0]
            if add_class not in classes:
                classes = classes + '\n' + add_class
        document.paragraphs[1].add_run(classes)
        for i in range(1, len(questions_list) - 1):
            self.fill_table(i + 1, document, i, report_holder)
            question_to_print = ''.join([j if ord(j) < 128 else ' ' for j in questions_list[i + 1]])
            document.paragraphs[i + 1].add_run(question_to_print)
        name = report_holder[0][1] + '.docx'
        if name == '#N/A.docx':
            name = report_holder[0][0] + '.docx'
            document.save('reports/' + name)
        else:
            document.save('reports/' + name)

    def calc_average_for_array(self, array):
        holder = 0
        for number in array:
            holder = holder + number
        average = holder / len(array)
        return average

    def print_master_report(self, reports, convert_table):
        q1_average = []
        q2_average = []
        q4_average = []
        q5_average = []
        q3_average = []
        document = Document('Master Template.docx')
        document.paragraphs[1].add_run(str(len(reports)))

        for part in reports:
            for entry in part:
                q1_average.append(entry[2])
                q2_average.append(entry[3])
                q3_average.append(entry[4])
                q4_average.append(entry[5])
                q5_average.append(entry[6])
        q1_average = self.calc_average_for_array(self.convert_to_number(q1_average, convert_table))
        q2_average = self.calc_average_for_array(self.convert_to_number(q2_average, convert_table))
        q3_average = self.calc_average_for_array(self.convert_to_number(q3_average, convert_table))
        q4_average = self.calc_average_for_array(self.convert_to_number(q4_average, convert_table))
        q5_average = self.calc_average_for_array(self.convert_to_number(q5_average, convert_table))
        document.paragraphs[3].add_run(str(q1_average))
        document.paragraphs[5].add_run(str(q2_average))
        document.paragraphs[7].add_run(str(q3_average))
        document.paragraphs[9].add_run(str(q4_average))
        document.paragraphs[11].add_run(str(q5_average))
        document.save('Master Report.docx')

    def run(self):
        with open("SP16 TA LA Eval.csv", "r") as readfile:
            questions = []
            response_holder = []
            report_list = []
            avoid_table = self.avoid_table_construct()
            number_table = self.conversion_table_construct()
            counter = 0

            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1

            for learning_assistant in self.create_name_list(response_holder):
                if learning_assistant == self.LaName:
                    report = self.create_individual_report(response_holder, learning_assistant)
                    report_list.append(report)
                    self.print_individual_report(report, questions)
            #self.print_master_report(report_list, number_table)
        readfile.close()

class Generate_Master_Report(QThread):
    def __init__(self,filepath,Template):
        QThread.__init__(self)
        self.filepath = filepath
        self.template = Template
    def __del__(self):
        self.wait()

    def conversion_table_construct(self):
        marker = 0
        conversion_table = {}
        grade = 0
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    marker = 0
                elif 'Eval Words' in row:
                    marker = 1
                elif 'Ignore Words' in row:
                    marker = 0
                elif marker is 1 and row is not "\n" or '':
                    grade += 1
                    conversion_table[row[:-1]] = grade
        textids.close()
        print "looking for these words"
        print conversion_table
        return conversion_table

    def avoid_table_construct(self):
        marker = 0
        avoid_list = []
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Ignore Words' in row:
                    marker = 1
                elif marker is 1 and row is not "\n":
                    avoid_list.append(row[:-1])
        textids.close()
        print "avoiding these words"
        print avoid_list
        return avoid_list

    def convert_to_number(self,convert_text, conversion_table):
        number = 0
        number_array = []
        for text in convert_text:
            for key in conversion_table.keys():
                if key in text:
                    number = conversion_table[key]
                    number_array.append(number)
                elif number == 0 or number == 'Null':
                    number = 'Null'
        return number_array

    def trim_response(self,text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self,question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:

                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

    def response_find(self, name, question_to_search, responses):
        ans = []
        for entry in responses:
            if name in entry:
                ans.append(entry[question_to_search])
        return ans

    def create_individual_report(self, response_list, name):
        report_data = []
        for entry in response_list:
            if name in entry[1]:
                report_data.append(entry)
        return report_data

    def create_name_list(self, response_list):
        name_list = []
        for entry in response_list:
            if entry[1] in name_list:
                continue
            else:
                name_list.append(entry[1])
        return name_list

    def response_options(self, question_number, report_array):
        rating_list = []
        for entry in report_array:
            add_rating = entry[question_number]
            if add_rating not in rating_list:
                rating_list.append(add_rating)
        return rating_list

    def fill_table(self,question_number, document, table_number, report_container):
        rating_list = []
        for entry in report_container:
            add_rating = entry[question_number]
            rating_list.append(add_rating)
        rating_catagories = self.response_options(question_number, report_container)
        rating_catagories.sort()
        rating_list.sort()
        total = 0
        new_row = document.tables[table_number]
        toi = document.tables[table_number]

        for rating in rating_catagories:
            new_row.add_row()
            toi.cell(-1, 0).text = rating
            for entry in rating_list:
                if entry == rating:
                    total += 1
                else:
                    total = total
            getcontext().prec = 2
            average = Decimal(total / int(document.tables[0].cell(1, 1).text))
            toi.cell(-1, 1).text = str(total)
            toi.cell(-1, 2).text = str(average * 100) + '%'
            total = 0

    def print_individual_report(self,report_holder, questions_list):
        document = Document('Individual Template.docx')
        document.paragraphs[0].add_run(report_holder[0][1])
        document.tables[0].cell(1, 1).text = str(len(report_holder))
        classes = ''

        for entry in report_holder:
            add_class = entry[0]
            if add_class not in classes:
                classes = classes + '\n' + add_class
        document.paragraphs[1].add_run(classes)
        for i in range(1, len(questions_list) - 1):
            self.fill_table(i + 1, document, i, report_holder)
            question_to_print = ''.join([j if ord(j) < 128 else ' ' for j in questions_list[i + 1]])
            document.paragraphs[i + 1].add_run(question_to_print)
        name = report_holder[0][1] + '.docx'
        if name == '#N/A.docx':
            name = report_holder[0][0] + '.docx'
            document.save('reports/' + name)
        else:
            document.save('reports/' + name)

    def calc_average_for_array(self,array):
        holder = 0
        for number in array:
            holder = holder + number
        average = holder / len(array)
        return average

    def print_master_report(self, reports, convert_table):
        q1_average = []
        q2_average = []
        q4_average = []
        q5_average = []
        q3_average = []
        document = Document('Master Template.docx')
        document.paragraphs[1].add_run(str(len(reports)))

        for part in reports:
            for entry in part:
                q1_average.append(entry[2])
                q2_average.append(entry[3])
                q3_average.append(entry[4])
                q4_average.append(entry[5])
                q5_average.append(entry[6])
        q1_average = self.calc_average_for_array(self.convert_to_number(q1_average, convert_table))
        q2_average = self.calc_average_for_array(self.convert_to_number(q2_average, convert_table))
        q3_average = self.calc_average_for_array(self.convert_to_number(q3_average, convert_table))
        q4_average = self.calc_average_for_array(self.convert_to_number(q4_average, convert_table))
        q5_average = self.calc_average_for_array(self.convert_to_number(q5_average, convert_table))
        document.paragraphs[3].add_run(str(q1_average))
        document.paragraphs[5].add_run(str(q2_average))
        document.paragraphs[7].add_run(str(q3_average))
        document.paragraphs[9].add_run(str(q4_average))
        document.paragraphs[11].add_run(str(q5_average))
        document.save('Master Report.docx')

    def run(self):
        with open("SP16 TA LA Eval.csv", "r") as readfile:
            questions = []
            response_holder = []
            report_list = []
            avoid_table = self.avoid_table_construct()
            number_table = self.conversion_table_construct()
            counter = 0

            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1

            for learning_assistant in self.create_name_list(response_holder):
                if learning_assistant not in report_list:
                    report = self.create_individual_report(response_holder, learning_assistant)
                    report_list.append(report)
                    #self.print_individual_report(report, questions)
            self.print_master_report(report_list, number_table)
        readfile.close()

class Generate_All_Reports(QThread):
    def __init__(self,filepath,Template):
        QThread.__init__(self)
        self.filepath = filepath
        self.template = Template
    def __del__(self):
        self.wait()

    def conversion_table_construct(self):
        marker = 0
        conversion_table = {}
        grade = 0
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    marker = 0
                elif 'Eval Words' in row:
                    marker = 1
                elif 'Ignore Words' in row:
                    marker = 0
                elif marker is 1 and row is not "\n" or '':
                    grade += 1
                    conversion_table[row[:-1]] = grade
        textids.close()
        print "looking for these words"
        print conversion_table
        return conversion_table

    def avoid_table_construct(self):
        marker = 0
        avoid_list = []
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Ignore Words' in row:
                    marker = 1
                elif marker is 1 and row is not "\n":
                    avoid_list.append(row[:-1])
        textids.close()
        print "avoiding these words"
        print avoid_list
        return avoid_list

    def convert_to_number(self,convert_text, conversion_table):
        number = 0
        number_array = []
        for text in convert_text:
            for key in conversion_table.keys():
                if key in text:
                    number = conversion_table[key]
                    number_array.append(number)
                elif number == 0 or number == 'Null':
                    number = 'Null'
        return number_array

    def trim_response(self,text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self,question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:

                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

    def response_find(self, name, question_to_search, responses):
        ans = []
        for entry in responses:
            if name in entry:
                ans.append(entry[question_to_search])
        return ans

    def create_individual_report(self, response_list, name):
        report_data = []
        for entry in response_list:
            if name in entry[1]:
                report_data.append(entry)
        return report_data

    def create_name_list(self, response_list):
        name_list = []
        for entry in response_list:
            if entry[1] in name_list:
                continue
            else:
                name_list.append(entry[1])
        return name_list

    def response_options(self, question_number, report_array):
        rating_list = []
        for entry in report_array:
            add_rating = entry[question_number]
            if add_rating not in rating_list:
                rating_list.append(add_rating)
        return rating_list

    def fill_table(self,question_number, document, table_number, report_container):
        rating_list = []
        for entry in report_container:
            add_rating = entry[question_number]
            rating_list.append(add_rating)
        rating_catagories = self.response_options(question_number, report_container)
        rating_catagories.sort()
        rating_list.sort()
        total = 0
        new_row = document.tables[table_number]
        toi = document.tables[table_number]

        for rating in rating_catagories:
            new_row.add_row()
            toi.cell(-1, 0).text = rating
            for entry in rating_list:
                if entry == rating:
                    total += 1
                else:
                    total = total
            getcontext().prec = 2
            average = Decimal(total / int(document.tables[0].cell(1, 1).text))
            toi.cell(-1, 1).text = str(total)
            toi.cell(-1, 2).text = str(average * 100) + '%'
            total = 0

    def print_individual_report(self,report_holder, questions_list):
        document = Document('Individual Template.docx')
        document.paragraphs[0].add_run(report_holder[0][1])
        document.tables[0].cell(1, 1).text = str(len(report_holder))
        classes = ''

        for entry in report_holder:
            add_class = entry[0]
            if add_class not in classes:
                classes = classes + '\n' + add_class
        document.paragraphs[1].add_run(classes)
        for i in range(1, len(questions_list) - 1):
            self.fill_table(i + 1, document, i, report_holder)
            question_to_print = ''.join([j if ord(j) < 128 else ' ' for j in questions_list[i + 1]])
            document.paragraphs[i + 1].add_run(question_to_print)
        name = report_holder[0][1] + '.docx'
        if name == '#N/A.docx':
            name = report_holder[0][0] + '.docx'
            document.save('reports/' + name)
        else:
            document.save('reports/' + name)

    def calc_average_for_array(self,array):
        holder = 0
        for number in array:
            holder = holder + number
        average = holder / len(array)
        return average

    def print_master_report(self, reports, convert_table):
        q1_average = []
        q2_average = []
        q4_average = []
        q5_average = []
        q3_average = []
        document = Document('Master Template.docx')
        document.paragraphs[1].add_run(str(len(reports)))

        for part in reports:
            for entry in part:
                q1_average.append(entry[2])
                q2_average.append(entry[3])
                q3_average.append(entry[4])
                q4_average.append(entry[5])
                q5_average.append(entry[6])
        q1_average = self.calc_average_for_array(self.convert_to_number(q1_average, convert_table))
        q2_average = self.calc_average_for_array(self.convert_to_number(q2_average, convert_table))
        q3_average = self.calc_average_for_array(self.convert_to_number(q3_average, convert_table))
        q4_average = self.calc_average_for_array(self.convert_to_number(q4_average, convert_table))
        q5_average = self.calc_average_for_array(self.convert_to_number(q5_average, convert_table))
        document.paragraphs[3].add_run(str(q1_average))
        document.paragraphs[5].add_run(str(q2_average))
        document.paragraphs[7].add_run(str(q3_average))
        document.paragraphs[9].add_run(str(q4_average))
        document.paragraphs[11].add_run(str(q5_average))
        document.save('Master Report.docx')

    def run(self):
        with open("SP16 TA LA Eval.csv", "r") as readfile:
            questions = []
            response_holder = []
            report_list = []
            avoid_table = self.avoid_table_construct()
            number_table = self.conversion_table_construct()
            counter = 0

            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1

            for learning_assistant in self.create_name_list(response_holder):
                if learning_assistant not in report_list:
                    report = self.create_individual_report(response_holder, learning_assistant)
                    report_list.append(report)
                    self.print_individual_report(report, questions)
            self.print_master_report(report_list, number_table)
        readfile.close()

class fillNames(QThread):
    def __init__(self,filepath):
        QThread.__init__(self)
        self.filepath = filepath

    def __del__(self):
        self.wait()

    def run(self):
        with open(self.filepath, "r") as readfile:
            self.questions = []
            response_holder = []
            # avoid_table = avoid_table_construct()
            # number_table = conversion_table_construct()
            counter = 0
            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    self.questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1
            for learning_assistant in self.create_name_list(response_holder):
                self.emit(SIGNAL('addLaName(QString)'), learning_assistant)
        readfile.close()

    def create_name_list(self, response_list):
        name_list = []
        for response in response_list:
            if response[1] in name_list:
                continue
            else:
                name_list.append(response[1])
        return name_list

    def trim_response(self, text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self, question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:
                    print 'true'
                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

class Create_Qualtrics_CSV(QThread):
    def __init__(self, filename):
        QThread.__init__(self)
        self.filepath = filename

    def __del__(self):
        self.wait()

    def run(self):
        fullNameList, courseList =self.generateLists(self.filepath)
        self.createCSV(fullNameList,courseList)



    def generateLists(self,filename):
        log = open(filename, "rb")
        fullNameList = []
        courseList = []
        reader = csv.DictReader(log)
        for row in reader:
            print "test"
            firstName = row["LA/TA_FIRST"]
            lastName = row["LA/TA_LAST"]
            subCode = row["Subject Code"]
            catalogNum = row["Catalog Number"]
            classCode = row["Class Section Code"]

            fullName = str(firstName) + " " + str(lastName)
            fullNameList.append(fullName)

            if len(str(classCode)) > 1:
                classCode = "00" + str(classCode)
            course = str(subCode) + " " + str(catalogNum) + "-" + classCode
            courseList.append(course)

        return fullNameList, courseList

    def createCSV(self,fullNamelist, courseList):
        output = open("Insert into Qualtrics.csv", "w")
        writer = csv.DictWriter(output, fieldnames=["Course Listing", "Full Name"], lineterminator='\n')
        writer.writeheader()
        for name, course in zip(fullNamelist, courseList):
            writer.writerow({"Full Name": name, "Course Listing": course})

def main():
    app = QtGui.QApplication(sys.argv)
    form = Main_App()
    form.show()
    app.exec_()

if __name__ == '__main__':
    main()