import os, tkFileDialog, csv, difflib
from Tkinter import Tk, BOTH
from ttk import Frame, Button, Style, Label, Entry
from csv import *
from collections import defaultdict
from docx import  Document
class GUI(Frame):

    def __init__(self,parent):
        Frame.__init__(self, parent)
        self.parent = parent
        self.initUI()

    def initUI(self):
        self.parent.title("Penn State University: LA Evaluation Report Generator")
        self.grid()
        self.style = Style()
        self.style.theme_use("alt")
        self.pack(fill =BOTH, expand = 1)

        key_label = Label(self, text='Specific LA: ')
        key_label.place(x=50,y=150)

        key = Entry(self)
        key.place(x=115, y=150)

        load_file_button = Button(self, text='Create Qualtrics CSV',command = all_reports)
        load_file_button.place(x=140,y=50)

        submit_key = Button(self, text='Single Report', command=lambda: specific_user(key.get()))
        submit_key.place(x=250, y=145)

        master_report = Button(self, text='Master Report',command = all_reports)
        master_report.place(x=85,y=100)

        all_report_button2 = Button(self, text='All Individual Reports',command = all_reports)
        all_report_button2.place(x=185,y=100)


def generateLists():
    filename = tkFileDialog.askopenfilename()
    log = open(filename, "rb")
    fullNameList = []
    courseList = []
    reader = csv.DictReader(log)
    for row in reader:
        firstName = row["LA/TA_FIRST"]
        lastName = row["LA/TA_LAST"]
        subCode = row["Subject Code"]
        catalogNum = row["Catalog Number"]
        classCode = row["Class Section Code"]

        fullName = str(firstName) + " " + str(lastName)
        fullNameList.append(fullName)

        if len(str(classCode)) > 1:
            classCode = "00" + str(classCode)
        course = str(subCode) + " " + str(catalogNum) + "-" + classCode
        courseList.append(course)

    return fullNameList, courseList
def strip_question(strip_text):
    word = strip_text.split('<', 1)[-1]
    word = word.split('>', 1)[-2]
    return word


def strip_answer(strip_text):
    word = strip_text.split('>', 1)[-1]
    word = word.split('<', 1)[-2]
    return word


def conversion_table_construct():
    marker = 0
    conversion_table = {}
    grade = 0
    with open('checkwords.txt') as textids:
        for row in textids:
            if '##' in row:
                continue
            elif 'Eval Words' in row:
                marker = 1
            elif 'Ignore Words' in row:
                marker = 0
            elif marker is 1 and row is not "\n" or '':
                grade += 1
                conversion_table[row[:-1]] = grade
    textids.close()
    print "looking for these words"
    print conversion_table
    return conversion_table


def avoid_table_construct():
    marker = 0
    avoid_list = []
    with open('checkwords.txt') as textids:
        for row in textids:
            if '##' in row:
                continue
            elif 'Ignore Words' in row:
                marker = 1
            elif marker is 1 and row is not "\n":
                avoid_list.append(row[:-1])
    textids.close()
    print "avoiding these words"
    print avoid_list
    return avoid_list


def convert_to_number(convert_text, conversion_table):
    number = 0
    for key in conversion_table.keys():
        if key in convert_text:
            number = conversion_table[key]
    if number == 0:
        number = convert_text
    """for entry in avoid_table:
        if entry in line:
            number = 0"""
    return number


def add_to_array(quest_array, resp_array, index_point, storage_array):
    if quest_array not in storage_array:
        storage_array[index_point].append(quest_array)
        storage_array[index_point].append(resp_array)
    else:
        storage_array[index_point].append(resp_array)


def create_response(storage_array, sorting_array):
    for group in storage_array:
        key = storage_array[group][1] + " " + storage_array[group][3]
        counter = 1
        for entry in storage_array[group]:
            if counter % 2 is 0:
                sorting_array[key].append(entry)
                counter += 1
            else:
                counter += 1
    return sorting_array


def create_print_list(storage_array, response_format_array):
    temp_responses = []
    for counter in range(0, storage_array.__len__()):
        for second_count in range(0, storage_array[0].__len__(), 2):
            holder = storage_array[counter][second_count] + " " + str(storage_array[counter][second_count + 1])
            temp_responses.append(holder)
    response_format_array.append(temp_responses)
    return response_format_array


def print_individual(learning_ass, profile_array):
    document = Document()
    paragraph = document.add_paragraph()
    for portion in profile_array[learning_ass]:
        paragraph.add_run(portion)
        paragraph.add_run('\n')
    alpha = profile_array[learning_ass][1] + '.docx'
    document.save(alpha)


def print_to_doc(profile_array):
    document = Document()
    for key in profile_array.keys():
        paragraph = document.add_paragraph()
        for bit in profile_array[key]:
            paragraph.add_run(bit)
            paragraph.add_run('\n')
    alpha = 'MasterList.docx'
    document.save(alpha)


def create_searchtext(sorting_array, searchtype):
    search_text = []
    search_course = []
    search_student = []
    for key in sorting_array.keys():
        course = key.split(' ', 1)[0]
        student = key.split(' ', 1)[1]
        if course not in search_course:
            search_course.append(course)
        else:
            continue
        if student not in search_student:
            search_student.append(student)
        else:
            continue
    search_text.append(search_course)
    search_text.append(search_student)
    if searchtype is 0:
        return search_course
    elif searchtype is 1:
        return search_student
    else:
        return search_text




def createCSV(fullNamelist, courseList):
    output = open("Insert into Qualtrics.csv", "w")
    writer = csv.DictWriter(output, fieldnames=["Course Listing", "Full Name"],lineterminator='\n')
    writer.writeheader()
    for name,course in zip(fullNamelist,courseList):
        writer.writerow({"Full Name" : name, "Course Listing": course})


def single_report(name):
    print "generated %s"%name

def all_reports():

    fullName,course = generateLists()
    createCSV(fullName, course)

def specific_user(name):
    return name

def xmlStuff():
    with open("TALA_Test_Eval.xml", "r") as readfile:
        checkText = defaultdict(list)
        indexNum = 0
        responses = []
        response_holder = []
        profile = defaultdict(list)
        avoid_table = avoid_table_construct()
        number_table = conversion_table_construct()
        for line in readfile:
            if '<Q' in line:
                question = strip_question(line)
                response = strip_answer(line)
                average_point = convert_to_number(response, number_table)
                add_to_array(question, response, indexNum, checkText)
            elif '</Response>' in line:
                indexNum += 1
        create_response(checkText, profile)
    readfile.close()
    for person in profile:
        print_individual(person, profile)
    print_to_doc(profile)
def main():
    root = Tk()
    root.geometry("400x250+300+300")
    app = GUI(root)
    root.mainloop()
    xmlStuff()







if __name__ == '__main__':
    main()