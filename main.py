from PyQt4 import QtGui
from PyQt4.Qt import *
from PyQt4.QtCore import QThread
from csv import *
from collections import defaultdict
from docx import Document
from PyQt4 import QtCore, QtGui
import sys, os, csv, difflib,random
import design


class Main_App(QtGui.QMainWindow, design.Ui_MainWindow):

    def __init__(self,parent=None):
        super(self.__class__, self).__init__()
        self.setupUi(self)

        self.filepath = ''

        self.qualtricsButton.clicked.connect(self.createQuatrics)
        self.loadFileButton.clicked.connect(self.loadFile)
        self.allReportsButton.clicked.connect(self.allReports)
        self.singleReportButton.clicked.connect(self.singleReport)

    def allReports(self):
        filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(filepath) > 0:
            self.myThread = Generate_All_Reports(filepath)
            self.connect(self.myThread, SIGNAL("finished()"), self.Done)
            self.myThread.start()
    def createQuatrics(self):
        filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(filepath) > 0:
            self.myThread = Create_Qualtrics_CSV(filepath)
            self.connect(self.myThread, SIGNAL("finished()"), self.QualDone())
            self.myThread.start()
    def QualDone(self):
        QtGui.QMessageBox.information(self, "Done!", "Qualtrics CSV Created in Current Working Directory!")

    def Done(self):
        QtGui.QMessageBox.information(self, "Done!", "All TA/LA reports generated in Current Working Directory!")

    def loadFile(self):
        self.filepath = QtGui.QFileDialog.getOpenFileName(self)
        if len(self.filepath) > 0:
            self.myThread = fillNames(self.filepath)
            self.connect(self.myThread,SIGNAL('addLaName(QString)'), self.addName)
            self.myThread.start()

    def addName(self,name):
        self.listWidget.addItem(name)

    def singleReport(self):
        currentPos = self.listWidget.currentRow()
        item = self.listWidget.item(currentPos)
        LaName = item.text()
        if len(self.filepath) > 0:
            self.myThread = generateIndividualReport(self.filepath,LaName)
            self.myThread.start()
        else:
            QtGui.QMessageBox.information(self, "Error", "TA/LA CSV Not Loaded Yet! Please load to proceed")



class generateIndividualReport(QThread):
    def __init__(self, filepath, La_Name):
        QThread.__init__(self)
        self.filepath = filepath
        self.LaName = La_Name
    def __del__(self):
        self.wait()

    def run(self):
        with open(self.filepath, "r") as readfile:
            self.questions = []
            response_holder = []
            # avoid_table = avoid_table_construct()
            # number_table = conversion_table_construct()
            counter = 0
            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    self.questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1
            for learning_assistant in self.create_name_list(response_holder):
                report = self.create_individual_report(self.questions, response_holder, learning_assistant)
                self.print_individual_report(self.report_data)
        readfile.close()

    def strip_question(self, strip_text):
        word = strip_text.split('<', 1)[-1]
        word = word.split('>', 1)[-2]
        return word

    def strip_answer(self, strip_text):
        word = strip_text.split('>', 1)[-1]
        word = word.split('<', 1)[-2]
        return word

    def conversion_table_construct(self):
        marker = 0
        conversion_table = {}
        grade = 0
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Eval Words' in row:
                    marker = 1
                elif 'Ignore Words' in row:
                    marker = 0
                elif marker is 1 and row is not "\n" or '':
                    grade += 1
                    conversion_table[row[:-1]] = grade
        textids.close()
        print "looking for these words"
        print conversion_table
        return conversion_table

    def avoid_table_construct(self):
        marker = 0
        self.avoid_list = []
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Ignore Words' in row:
                    marker = 1
                elif marker is 1 and row is not "\n":
                    self.avoid_list.append(row[:-1])
        textids.close()
        print "avoiding these words"
        print self.avoid_list

    def convert_to_number(self, convert_text, conversion_table):
        number = 0
        for key in conversion_table.keys():
            if key in convert_text:
                number = conversion_table[key]
        if number == 0:
            number = convert_text
        """for entry in avoid_table:
            if entry in line:
                number = 0"""
        return number

    def add_to_array(self, quest_array, resp_array, index_point, storage_array):
        if quest_array not in storage_array:
            storage_array[index_point].append(quest_array)
            storage_array[index_point].append(resp_array)
        else:
            storage_array[index_point].append(resp_array)

    def create_response(self, storage_array, sorting_array):
        for group in storage_array:
            key = storage_array[group][1] + " " + storage_array[group][3]
            counter = 1
            for entry in storage_array[group]:
                if counter % 2 is 0:
                    sorting_array[key].append(entry)
                    counter += 1
                else:
                    counter += 1
        return sorting_array

    def create_print_list(self, storage_array, response_format_array):
        temp_responses = []
        for counter in range(0, storage_array.__len__()):
            for second_count in range(0, storage_array[0].__len__(), 2):
                holder = storage_array[counter][second_count] + " " + str(storage_array[counter][second_count + 1])
                temp_responses.append(holder)
        response_format_array.append(temp_responses)
        return response_format_array

    def calculate_tables_and_such(self, learning_ass, profile_array):
        number = random.randrange(1, 5)
        return number

    def print_individual(self, learning_ass, profile_array):
        document = Document()
        paragraph = document.add_paragraph()
        for portion in profile_array[learning_ass]:
            paragraph.add_run(portion)
            paragraph.add_run('\n')
        alpha = profile_array[learning_ass][1] + '.docx'
        document.save(alpha)

    def print_to_doc(self, profile_array):
        document = Document('template.docx')
        for key in profile_array.keys():
            paragraph = document.add_paragraph()
            """for bit in profile_array[key]:
                paragraph.add_run(bit)
                paragraph.add_run('\n')"""
            table = document.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Answer'
            hdr_cells[2].text = 'Bar'
            hdr_cells[3].text = 'Response'
            hdr_cells[4].text = '%'
            counter = 0
            for entry in profile_array[key]:
                if any(part in self.avoid_list for part in entry):
                    continue
                else:
                    row_cells = table.add_row().cells
                    counter += 1
                    row_cells[0].text = str(counter)
                    row_cells[1].text = 'question'
                    row_cells[2].text = entry
                    row_cells[3].text = 'null'
                    row_cells[4].text = 'null'

        alpha = 'Master.docx'
        document.save(alpha)

    def create_searchtext(self, sorting_array, searchtype):
        search_text = []
        search_course = []
        search_student = []
        for key in sorting_array.keys():
            course = key.split(' ', 1)[0]
            student = key.split(' ', 1)[1]
            if course not in search_course:
                search_course.append(course)
            else:
                continue
            if student not in search_student:
                search_student.append(student)
            else:
                continue
        search_text.append(search_course)
        search_text.append(search_student)
        if searchtype is 0:
            return search_course
        elif searchtype is 1:
            return search_student
        else:
            return search_text

    def trim_response(self, text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self, question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:
                    print 'true'
                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

    def question_responses(self, question_list, response_list):

        return "hi"

    def pair_off_indices(self, question_list, response_list):

        return "hi"

    def response_find(self, name, question_to_search, responses):
        ans = []
        for response in responses:
            if name in response:
                ans.append(response[question_to_search])
        return ans

    def create_individual_report(self, question_list, response_list, name):
        self.report_data = []
        for i in range(0, len(question_list) - 1):
            self.report_data.append([])
        for entry in response_list:
            if name in entry[1]:
                for i in range(0, len(question_list) - 1):
                    self.report_data[i].append(self.response_find(name, i, response_list))

    def create_name_list(self, response_list):
        name_list = []
        name_list.append(str(self.LaName))
        return name_list

    def fill_table(self, question_number, document, table_number):
        counter = 1
        rating_catagories = []
        rating_list = []
        for entry in self.report_data[question_number]:
            add_rating = entry[0]
            rating_list.append(add_rating)
            if add_rating not in rating_catagories:
                rating_catagories.append(add_rating)
        rating_catagories.sort()
        rating_list.sort()

        total = 0
        for rating in rating_catagories:
            document.tables[table_number].cell(counter, 0).text = rating
            for entry in rating_list:
                if entry is rating:
                    total += 1
                else:
                    continue
            document.tables[table_number].cell(counter, 1).text = str(total)
            document.tables[table_number].cell(counter, 2).text = str((total / len(rating_list) * 100))
            total = 0
            counter += 1

    def print_individual_report(self, report):
        document = Document('Template.docx')
        document.paragraphs[5].add_run(report[1][0])
        document.tables[0].cell(1, 1).text = str(len(report[1]))
        classes = ''
        for entry in report[0]:
            add_class = entry[0]
            if add_class not in classes:
                classes = classes + add_class
        document.paragraphs[7].add_run(classes)
        self.fill_table(2, document, 1)
        document.paragraphs[8].add_run(self.questions[2])
        self.fill_table(3, document, 2)
        document.paragraphs[9].add_run(self.questions[3])
        self.fill_table(4, document, 3)
        document.paragraphs[10].add_run(self.questions[4])
        self.fill_table(5, document, 4)
        document.paragraphs[11].add_run(self.questions[5])
        self.fill_table(6, document, 5)
        document.paragraphs[12].add_run(self.questions[6])
        self.fill_table(7, document, 6)
        document.paragraphs[13].add_run(self.questions[7])
        self.fill_table(8, document, 7)
        name = report[1][0][0] + '.docx'
        document.save(name)

class fillNames(QThread):
    def __init__(self,filepath):
        QThread.__init__(self)
        self.filepath = filepath

    def __del__(self):
        self.wait()

    def run(self):
        with open(self.filepath, "r") as readfile:
            self.questions = []
            response_holder = []
            # avoid_table = avoid_table_construct()
            # number_table = conversion_table_construct()
            counter = 0
            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    self.questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1
            for learning_assistant in self.create_name_list(response_holder):
                self.emit(SIGNAL('addLaName(QString)'), learning_assistant)
        readfile.close()

    def create_name_list(self, response_list):
        name_list = []
        for response in response_list:
            if response[1] in name_list:
                continue
            else:
                name_list.append(response[1])
        return name_list

    def trim_response(self, text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self, question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:
                    print 'true'
                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

class Generate_All_Reports(QThread):
    def __init__(self,filepath):
        QThread.__init__(self)
        self.filepath = filepath

    def __del__(self):
        self.wait()

    def run(self):
        with open(self.filepath, "r") as readfile:
            self.questions = []
            response_holder = []
            #avoid_table = avoid_table_construct()
            #number_table = conversion_table_construct()
            counter = 0
            for line in readfile:
                response = line.split(',')
                if counter is 0:
                    counter += 1
                elif counter is 1:
                    self.questions = self.trim_question(response)
                    counter += 1
                else:
                    response_holder.append(self.trim_response(response))
                    counter += 1
            for learning_assistant in self.create_name_list(response_holder):
                report = self.create_individual_report(self.questions, response_holder, learning_assistant)
                self.print_individual_report(self.report_data)
        readfile.close()




    def strip_question(self,strip_text):
        word = strip_text.split('<', 1)[-1]
        word = word.split('>', 1)[-2]
        return word

    def strip_answer(self,strip_text):
        word = strip_text.split('>', 1)[-1]
        word = word.split('<', 1)[-2]
        return word

    def conversion_table_construct(self):
        marker = 0
        conversion_table = {}
        grade = 0
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Eval Words' in row:
                    marker = 1
                elif 'Ignore Words' in row:
                    marker = 0
                elif marker is 1 and row is not "\n" or '':
                    grade += 1
                    conversion_table[row[:-1]] = grade
        textids.close()
        print "looking for these words"
        print conversion_table
        return conversion_table

    def avoid_table_construct(self):
        marker = 0
        self.avoid_list = []
        with open('checkwords.txt') as textids:
            for row in textids:
                if '##' in row:
                    continue
                elif 'Ignore Words' in row:
                    marker = 1
                elif marker is 1 and row is not "\n":
                    self.avoid_list.append(row[:-1])
        textids.close()
        print "avoiding these words"
        print self.avoid_list

    def convert_to_number(self,convert_text, conversion_table):
        number = 0
        for key in conversion_table.keys():
            if key in convert_text:
                number = conversion_table[key]
        if number == 0:
            number = convert_text
        """for entry in avoid_table:
            if entry in line:
                number = 0"""
        return number

    def add_to_array(self,quest_array, resp_array, index_point, storage_array):
        if quest_array not in storage_array:
            storage_array[index_point].append(quest_array)
            storage_array[index_point].append(resp_array)
        else:
            storage_array[index_point].append(resp_array)

    def create_response(self,storage_array, sorting_array):
        for group in storage_array:
            key = storage_array[group][1] + " " + storage_array[group][3]
            counter = 1
            for entry in storage_array[group]:
                if counter % 2 is 0:
                    sorting_array[key].append(entry)
                    counter += 1
                else:
                    counter += 1
        return sorting_array

    def create_print_list(self,storage_array, response_format_array):
        temp_responses = []
        for counter in range(0, storage_array.__len__()):
            for second_count in range(0, storage_array[0].__len__(), 2):
                holder = storage_array[counter][second_count] + " " + str(storage_array[counter][second_count + 1])
                temp_responses.append(holder)
        response_format_array.append(temp_responses)
        return response_format_array

    def calculate_tables_and_such(self,learning_ass, profile_array):
        number = random.randrange(1, 5)
        return number

    def print_individual(self,learning_ass, profile_array):
        document = Document()
        paragraph = document.add_paragraph()
        for portion in profile_array[learning_ass]:
            paragraph.add_run(portion)
            paragraph.add_run('\n')
        alpha = profile_array[learning_ass][1] + '.docx'
        document.save(alpha)

    def print_to_doc(self,profile_array):
        document = Document('template.docx')
        for key in profile_array.keys():
            paragraph = document.add_paragraph()
            """for bit in profile_array[key]:
                paragraph.add_run(bit)
                paragraph.add_run('\n')"""
            table = document.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Answer'
            hdr_cells[2].text = 'Bar'
            hdr_cells[3].text = 'Response'
            hdr_cells[4].text = '%'
            counter = 0
            for entry in profile_array[key]:
                if any(part in self.avoid_list for part in entry):
                    continue
                else:
                    row_cells = table.add_row().cells
                    counter += 1
                    row_cells[0].text = str(counter)
                    row_cells[1].text = 'question'
                    row_cells[2].text = entry
                    row_cells[3].text = 'null'
                    row_cells[4].text = 'null'

        alpha = 'Master.docx'
        document.save(alpha)

    def create_searchtext(self,sorting_array, searchtype):
        search_text = []
        search_course = []
        search_student = []
        for key in sorting_array.keys():
            course = key.split(' ', 1)[0]
            student = key.split(' ', 1)[1]
            if course not in search_course:
                search_course.append(course)
            else:
                continue
            if student not in search_student:
                search_student.append(student)
            else:
                continue
        search_text.append(search_course)
        search_text.append(search_student)
        if searchtype is 0:
            return search_course
        elif searchtype is 1:
            return search_student
        else:
            return search_text

    def trim_response(self,text_to_trim):
        return_text = []
        for i in range(10, len(text_to_trim)):
            return_text.append(text_to_trim[i])
        return return_text

    def trim_question(self,question_to_trim):
        question_text = []
        for i in range(10, len(question_to_trim) - 4):
            if '${' in question_to_trim[i]:
                question_text_p1 = question_to_trim[i].split('${', 1)[-2]
                question_text_p2a = question_to_trim[i].split('${', 1)[-1]
                question_text_p2 = question_text_p2a.split('}', 1)[-1]
                question_to_append = question_text_p1 + ' ' + question_text_p2
                if '\xc2\xa0' in question_to_append:
                    question_text_p3 = question_to_append.split('\xc2\xa0')[0]
                    question_text_p4 = question_to_append.split('\xc2\xa0')[-1]
                    question_to_append = question_text_p3 + ' your TA/LA ' + question_text_p4
                elif 'Assistant\xe2\x80\x99s' in question_to_append:
                    print 'true'
                    question_text_p5 = question_to_append.split('Assistant\xe2\x80\x99s')[0]
                    question_text_p6 = question_to_append.split('Assistant\xe2\x80\x99s')[-1]
                    question_to_append = question_text_p5 + "Assistant's" + question_text_p6
                question_text.append(question_to_append)
            else:
                question_text.append(question_to_trim[i])
        return question_text

    def question_responses(self,question_list, response_list):

        return "hi"

    def pair_off_indices(self,question_list, response_list):

        return "hi"

    def response_find(self,name, question_to_search, responses):
        ans = []
        for response in responses:
            if name in response:
                ans.append(response[question_to_search])
        return ans

    def create_individual_report(self,question_list, response_list, name):
        self.report_data = []
        for i in range(0, len(question_list) - 1):
            self.report_data.append([])
        for entry in response_list:
            if name in entry[1]:
                for i in range(0, len(question_list) - 1):
                    self.report_data[i].append(self.response_find(name, i, response_list))


    def create_name_list(self,response_list):
        name_list = []
        for response in response_list:
            if response[1] in name_list:
                continue
            else:
                name_list.append(response[1])
        return name_list

    def fill_table(self,question_number, document, table_number):
        counter = 1
        rating_catagories = []
        rating_list = []
        for entry in self.report_data[question_number]:
            add_rating = entry[0]
            rating_list.append(add_rating)
            if add_rating not in rating_catagories:
                rating_catagories.append(add_rating)
        rating_catagories.sort()
        rating_list.sort()

        total = 0
        for rating in rating_catagories:
            document.tables[table_number].cell(counter, 0).text = rating
            for entry in rating_list:
                if entry is rating:
                    total += 1
                else:
                    continue
            document.tables[table_number].cell(counter, 1).text = str(total)
            document.tables[table_number].cell(counter, 2).text = str((total / len(rating_list) * 100))
            total = 0
            counter += 1

    def print_individual_report(self,report):
        document = Document('Template.docx')
        document.paragraphs[5].add_run(report[1][0])
        document.tables[0].cell(1, 1).text = str(len(report[1]))
        classes = ''
        for entry in report[0]:
            add_class = entry[0]
            if add_class not in classes:
                classes = classes + add_class
        document.paragraphs[7].add_run(classes)
        self.fill_table(2, document, 1)
        document.paragraphs[8].add_run(self.questions[2])
        self.fill_table(3, document, 2)
        document.paragraphs[9].add_run(self.questions[3])
        self.fill_table(4, document, 3)
        document.paragraphs[10].add_run(self.questions[4])
        self.fill_table(5, document, 4)
        document.paragraphs[11].add_run(self.questions[5])
        self.fill_table(6, document, 5)
        document.paragraphs[12].add_run(self.questions[6])
        self.fill_table(7, document, 6)
        document.paragraphs[13].add_run(self.questions[7])
        self.fill_table(8, document, 7)
        name = report[1][0][0] + '.docx'
        document.save(name)

class Create_Qualtrics_CSV(QThread):
    def __init__(self, filename):
        QThread.__init__(self)
        self.filepath = filename

    def __del__(self):
        self.wait()

    def run(self):
        fullNameList, courseList =self.generateLists(self.filepath)
        self.createCSV(fullNameList,courseList)



    def generateLists(self,filename):
        log = open(filename, "rb")
        fullNameList = []
        courseList = []
        reader = csv.DictReader(log)
        for row in reader:
            firstName = row["LA/TA_FIRST"]
            lastName = row["LA/TA_LAST"]
            subCode = row["Subject Code"]
            catalogNum = row["Catalog Number"]
            classCode = row["Class Section Code"]

            fullName = str(firstName) + " " + str(lastName)
            fullNameList.append(fullName)

            if len(str(classCode)) > 1:
                classCode = "00" + str(classCode)
            course = str(subCode) + " " + str(catalogNum) + "-" + classCode
            courseList.append(course)

        return fullNameList, courseList

    def createCSV(self,fullNamelist, courseList):
        output = open("Insert into Qualtrics.csv", "w")
        writer = csv.DictWriter(output, fieldnames=["Course Listing", "Full Name"], lineterminator='\n')
        writer.writeheader()
        for name, course in zip(fullNamelist, courseList):
            writer.writerow({"Full Name": name, "Course Listing": course})

def main():
    app = QtGui.QApplication(sys.argv)
    form = Main_App()
    form.show()
    app.exec_()

if __name__ == '__main__':
    main()